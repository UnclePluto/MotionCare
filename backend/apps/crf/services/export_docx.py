from pathlib import Path

from django.conf import settings
from docx import Document


def export_preview_to_docx(preview: dict, output_path: Path) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document(settings.CRF_TEMPLATE_PATH)
    document.add_page_break()
    document.add_heading("系统生成数据摘要", level=1)
    document.add_paragraph(f"患者姓名：{preview['patient']['name']}")
    document.add_paragraph(f"项目名称：{preview['project']['name']}")
    document.add_paragraph(f"分组：{preview['group']['name']}")
    if preview["missing_fields"]:
        document.add_heading("缺失字段", level=2)
        for field in preview["missing_fields"]:
            document.add_paragraph(field, style="List Bullet")
    document.save(output_path)
    return output_path

